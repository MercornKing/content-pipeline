#!/usr/bin/env python3
"""
process.py — WorkSmart Reviews pipeline processor
Converts .docx files in /drafts to Markdown in /articles,
injects affiliate links, then auto-updates the ARTICLES array in
content/site.js so new articles go live without any manual editing.

Folder structure:
    drafts/           <- drop .docx files here
    articles/         <- processed .md files (served by Netlify)
    content/          <- site files: index.html, site.js, site.css
    affiliate_links.json
    topic_map.json    <- optional: maps slug keywords to topic keys

Usage:
    python process.py                  # process all new drafts
    python process.py --force          # reprocess all drafts
    python process.py --file foo.docx  # process a single file
"""

import argparse
import json
import logging
import re
import sys
from datetime import datetime
from pathlib import Path

import docx
from markdownify import markdownify

# Paths
ROOT         = Path(__file__).parent
DRAFTS_DIR   = ROOT / "drafts"
ARTICLES_DIR = ROOT / "content" / "articles"
SITE_JS      = ROOT / "content" / "site.js"
LINKS_FILE   = ROOT / "affiliate_links.json"
TOPIC_MAP    = ROOT / "topic_map.json"
LOG_FILE     = ROOT / "pipeline.log"

# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger(__name__)

# Topic keyword defaults
DEFAULT_TOPIC_KEYWORDS = {
    "marketing":    ["email", "marketing", "mailchimp", "mailerlite", "brevo", "newsletter", "seo", "social"],
    "productivity": ["monday", "asana", "notion", "trello", "clickup", "project", "task", "productivity", "crm"],
    "coaching":     ["coach", "coaching", "scheduling", "calendly", "acuity", "tidycal", "booking"],
    "finance":      ["accounting", "invoic", "bookkeep", "payroll", "tax", "vat", "xero", "quickbooks", "sage", "sole trader"],
    "hr":           ["hr", "hiring", "recruit", "payslip", "employee", "staff", "onboard"],
}

def load_affiliate_links():
    if not LINKS_FILE.exists():
        log.warning("affiliate_links.json not found — no links injected.")
        return {}
    with LINKS_FILE.open(encoding="utf-8") as f:
        raw = json.load(f)
    links = {k.lower().strip(): v for k, v in raw.items() if not k.startswith("_")}
    log.info(f"Loaded {len(links)} affiliate link(s)")
    return links

def infer_topic(slug, title):
    keywords = dict(DEFAULT_TOPIC_KEYWORDS)
    if TOPIC_MAP.exists():
        with TOPIC_MAP.open(encoding="utf-8") as f:
            keywords.update(json.load(f))
    haystack = (slug + " " + title).lower()
    for topic, kws in keywords.items():
        if any(kw in haystack for kw in kws):
            return topic
    return "general"

def title_to_excerpt(first_para):
    clean = re.sub(r"[#*_`]", "", first_para).strip()
    if len(clean) <= 160:
        return clean
    return clean[:160].rsplit(" ", 1)[0] + "..."

def docx_to_markdown(docx_path):
    doc = docx.Document(str(docx_path))
    html_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            html_parts.append("<br>")
            continue
        style = para.style.name.lower()
        if style.startswith("heading 1"):
            html_parts.append(f"<h1>{text}</h1>")
        elif style.startswith("heading 2"):
            html_parts.append(f"<h2>{text}</h2>")
        elif style.startswith("heading 3"):
            html_parts.append(f"<h3>{text}</h3>")
        else:
            line = ""
            for run in para.runs:
                t = run.text
                if run.bold and run.italic:
                    t = f"<strong><em>{t}</em></strong>"
                elif run.bold:
                    t = f"<strong>{t}</strong>"
                elif run.italic:
                    t = f"<em>{t}</em>"
                line += t
            html_parts.append(f"<p>{line}</p>")
    for table in doc.tables:
        html_parts.append("<table>")
        for i, row in enumerate(table.rows):
            html_parts.append("<tr>")
            tag = "th" if i == 0 else "td"
            for cell in row.cells:
                html_parts.append(f"<{tag}>{cell.text.strip()}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")
    raw_html = "\n".join(html_parts)
    md = markdownify(raw_html, heading_style="ATX", bullets="-")
    return re.sub(r"\n{3,}", "\n\n", md).strip()

def inject_affiliate_links(markdown, links):
    """
    Replace every 'Product Name [AFFILIATE LINK]' with a Markdown hyperlink.

    Three-stage matching (most to least specific):
      1. Exact lowercase match on the product name
      2. Strip trailing markdown/punctuation characters, try again
      3. Longest key that is a substring of the product name

    Unmatched placeholders are left unchanged and logged as warnings.
    Updated from original single-pass line-context approach to improve
    match rate, particularly for product names with punctuation variations.
    """
    pattern = re.compile(r'([^\[\n]{2,120})\s*\[AFFILIATE LINK\]', re.IGNORECASE)
    warnings = []
    count = 0

    def replace(m):
        nonlocal count
        name = m.group(1).strip()
        key  = name.lower()

        # Stage 1: exact match
        url = links.get(key)

        # Stage 2: strip trailing punctuation/markdown
        if not url:
            cleaned = re.sub(r'[\*_`#>\-]+$', '', key).strip()
            url = links.get(cleaned)

        # Stage 3: longest substring match
        if not url:
            candidates = [k for k in links if k in key]
            if candidates:
                url = links[max(candidates, key=len)]

        if url:
            count += 1
            log.info(f"  Injected link for '{name}'")
            return f"[{name}]({url})"
        else:
            msg = f"No affiliate link matched for: '{name}'"
            warnings.append(msg)
            log.warning(f"  {msg}")
            return m.group(0)

    result = pattern.sub(replace, markdown)
    return result, warnings

def add_frontmatter(markdown, source_filename, topic, excerpt, date):
    title_match = re.search(r"^#\s+(.+)$", markdown, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else source_filename
    safe_excerpt = excerpt.replace('"', '\\"')
    fm = f'---\ntitle: "{title}"\ndate: {date}\ntopic: {topic}\nexcerpt: "{safe_excerpt}"\nsource: "{source_filename}"\n---\n\n'
    return fm + markdown

def update_site_js(slug, title, topic, excerpt, date):
    if not SITE_JS.exists():
        log.warning(f"  site.js not found — skipping index update")
        return False
    js_text = SITE_JS.read_text(encoding="utf-8")

    if f"slug: '{slug}'" in js_text:
        log.info(f"  site.js already has '{slug}' — skipping")
        return True

    safe_title   = title.replace("'", "\\'")
    safe_excerpt = excerpt.replace('"', '\\"')
    new_entry = (
        f"  {{\n"
        f"    slug: '{slug}',\n"
        f"    title: '{safe_title}',\n"
        f"    topic: '{topic}',\n"
        f"    excerpt: \"{safe_excerpt}\",\n"
        f"    date: '{date}',\n"
        f"  }}"
    )

    updated = re.sub(
        r"(const ARTICLES\s*=\s*\[)([\s\S]*?)(\];)",
        lambda m: m.group(1) + m.group(2).rstrip(",\n ") + ",\n" + new_entry + "\n];",
        js_text
    )

    if updated == js_text:
        log.warning("  Could not find ARTICLES array in site.js")
        return False

    # Keep articleCount in sync if it exists in the file
    count = len(re.findall(r"slug:", updated))
    updated = re.sub(r"(articleCount:\s*)\d+", f"\\g<1>{count}", updated)

    SITE_JS.write_text(updated, encoding="utf-8")
    log.info(f"  Updated site.js with '{slug}', articleCount -> {count}")
    return True

def title_to_slug(title):
    """
    Convert an article H1 title to a short, stable URL slug.
    Strips em-dash subtitles and trailing years so the slug matches
    hand-curated slugs already in site.js.
    Examples:
      'Monday.com Review for Small Business (2026)'  -> 'monday-com-review-for-small-business'
      'Best Email Tools — UK Guide'                  -> 'best-email-tools'
    """
    base = re.split(r'\s*[—–]\s*|\s+-\s+', title)[0].strip()
    base = re.sub(r'\s+\d{4}$', '', base).strip()
    base = re.sub(r'\s*\(.*?\)$', '', base).strip()  # strip trailing parentheticals
    return re.sub(r'[^a-z0-9]+', '-', base.lower()).strip('-')


def process_file(docx_path, links, force=False):
    # Derive slug from H1 title so it matches hand-curated slugs in site.js.
    _slug_fallback = re.sub(r"[^a-z0-9]+", "-", docx_path.stem.lower()).strip("-")
    try:
        _preview_md  = docx_to_markdown(docx_path)
        _title_match = re.search(r"^#\s+(.+)$", _preview_md, re.MULTILINE)
        slug = title_to_slug(_title_match.group(1)) if _title_match else _slug_fallback
    except Exception:
        slug = _slug_fallback
    out_path = ARTICLES_DIR / f"{slug}.md"

    if out_path.exists() and not force:
        log.info(f"Skipping '{docx_path.name}' — already processed")
        return True

    log.info(f"Processing: {docx_path.name}")
    try:
        markdown = docx_to_markdown(docx_path)
    except Exception as e:
        log.error(f"  Conversion failed: {e}")
        return False

    markdown, warnings = inject_affiliate_links(markdown, links)

    title_match = re.search(r"^#\s+(.+)$", markdown, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else slug.replace("-", " ").title()
    topic = infer_topic(slug, title)
    date  = datetime.now().strftime("%Y-%m-%d")
    paras = [l.strip() for l in markdown.split("\n") if l.strip() and not l.startswith("#")]
    excerpt = title_to_excerpt(paras[0] if paras else title)

    markdown = add_frontmatter(markdown, docx_path.name, topic, excerpt, date)
    ARTICLES_DIR.mkdir(parents=True, exist_ok=True)
    out_path.write_text(markdown, encoding="utf-8")
    log.info(f"  Written to content/articles/{slug}.md")

    update_site_js(slug, title, topic, excerpt, date)

    if warnings:
        log.warning(f"  {len(warnings)} unmatched [AFFILIATE LINK] placeholder(s)")
    return True

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--force", action="store_true")
    parser.add_argument("--file", type=str)
    args = parser.parse_args()

    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    ARTICLES_DIR.mkdir(parents=True, exist_ok=True)

    links = load_affiliate_links()

    files = [DRAFTS_DIR / args.file] if args.file else sorted(DRAFTS_DIR.glob("*.docx"))
    if not files:
        log.info("No .docx files in /drafts — nothing to do.")
        sys.exit(0)

    log.info(f"Found {len(files)} file(s)")
    results = [process_file(f, links, force=args.force) for f in files]
    ok = sum(results)
    log.info(f"Done — {ok} succeeded, {len(results)-ok} failed")
    if len(results) - ok:
        sys.exit(1)

if __name__ == "__main__":
    main()
